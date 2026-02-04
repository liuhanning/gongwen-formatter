# -*- coding: utf-8 -*-
"""
公文级 Word 文档自动化排版引擎
符合 GB/T 9704 党政机关公文格式规范
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ==================== 常量定义 ====================

# 页面设置参数
PAGE_WIDTH = Mm(210)       # A4 宽度
PAGE_HEIGHT = Mm(297)      # A4 高度
MARGIN_TOP = Mm(37)        # 上边距
MARGIN_BOTTOM = Mm(35)     # 下边距
MARGIN_LEFT = Mm(28)       # 左边距
MARGIN_RIGHT = Mm(26)      # 右边距
HEADER_DISTANCE = Cm(1.8)  # 页眉距边界
FOOTER_DISTANCE = Cm(1.8)  # 页脚距边界

# 字体映射（带兼容性退化）
FONT_MAP = {
    'fzxbsjt': '方正小标宋简体',     # 一级标题
    'heiti': '黑体',                 # 二级标题
    'kaiti_gb2312': '楷体_GB2312',   # 三级标题
    'fangsong_gb2312': '仿宋_GB2312', # 四级标题/正文
    'tnr': 'Times New Roman',        # 英文/数字
}

# 字体退化方案
FONT_FALLBACK = {
    '方正小标宋简体': ['FZXiaoBiaoSong-B05', '华文中宋', '宋体'],
    '楷体_GB2312': ['楷体', 'KaiTi', 'SimKai'],
    '仿宋_GB2312': ['仿宋', 'FangSong', 'SimFang'],
}

# 字号映射（磅值）
FONT_SIZE = {
    '初号': Pt(42),
    '小初': Pt(36),
    '一号': Pt(26),
    '小一': Pt(24),
    '二号': Pt(22),
    '小二': Pt(18),
    '三号': Pt(16),
    '小三': Pt(15),
    '四号': Pt(14),
    '小四': Pt(12),
    '五号': Pt(10.5),
    '小五': Pt(9),
}

# 样式定义
STYLES = {
    'level1': {  # 一级标题
        'cn_font': '方正小标宋简体',
        'en_font': 'Times New Roman',
        'font_size': '二号',
        'line_spacing': Pt(30),
        'space_before': 0.5,  # 行
        'space_after': 0.5,   # 行
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'first_line_indent': 0,
        'left_indent': 0,
    },
    'level2': {  # 二级标题
        'cn_font': '黑体',
        'en_font': 'Times New Roman',
        'font_size': '三号',
        'line_spacing': Pt(30),
        'space_before': 0.5,
        'space_after': 0.5,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'first_line_indent': 0,
        'left_indent': 2,  # 字符
    },
    'level3': {  # 三级标题
        'cn_font': '楷体_GB2312',
        'en_font': 'Times New Roman',
        'font_size': '三号',
        'line_spacing': Pt(30),
        'space_before': 0.5,
        'space_after': 0.5,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'first_line_indent': 0,
        'left_indent': 2,
    },
    'level4': {  # 四级标题
        'cn_font': '仿宋_GB2312',
        'en_font': 'Times New Roman',
        'font_size': '三号',
        'line_spacing': Pt(28),
        'space_before': 0.5,
        'space_after': 0.5,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'first_line_indent': 0,
        'left_indent': 2,
    },
    'level5': {  # 五级标题（扩展）
        'cn_font': '仿宋_GB2312',
        'en_font': 'Times New Roman',
        'font_size': '三号',
        'line_spacing': Pt(28),
        'space_before': 0.5,
        'space_after': 0.5,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'first_line_indent': 0,
        'left_indent': 2,
    },
    'level6': {  # 六级（圆圈序号，用于正文列表项）
        'cn_font': '仿宋_GB2312',
        'en_font': 'Times New Roman',
        'font_size': '三号',
        'line_spacing': Pt(28),
        'space_before': 0,
        'space_after': 0,
        'alignment': WD_ALIGN_PARAGRAPH.LEFT,
        'first_line_indent': 2,
        'left_indent': 0,
    },
    'body': {  # 正文
        'cn_font': '仿宋_GB2312',
        'en_font': 'Times New Roman',
        'font_size': '三号',
        'line_spacing': Pt(28),
        'space_before': 0,
        'space_after': 0,
        'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'first_line_indent': 2,  # 首行缩进2字符
        'left_indent': 0,
    },
    'table_title': {  # 表格标题
        'cn_font': '黑体',
        'en_font': 'Times New Roman',
        'font_size': '小四',
        'line_spacing': Pt(20),
        'space_before': 0.5,
        'space_after': 0.5,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'first_line_indent': 0,
        'left_indent': 0,
        'bold': False,
    },
    'table_header': {  # 表头
        'cn_font': '黑体',
        'en_font': 'Times New Roman',
        'font_size': '小四',
        'line_spacing': None,  # 单倍行距
        'space_before': 0,
        'space_after': 0,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'first_line_indent': 0,
        'left_indent': 0,
        'bold': False,
    },
    'table_content': {  # 表格内容
        'cn_font': '仿宋_GB2312',
        'en_font': 'Times New Roman',
        'font_size': '小四',
        'line_spacing': None,  # 单倍行距
        'space_before': 0,
        'space_after': 0,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'first_line_indent': 0,
        'left_indent': 0,
    },
    'figure_title': {  # 插图标题
        'cn_font': '黑体',
        'en_font': 'Times New Roman',
        'font_size': '小四',
        'line_spacing': Pt(20),
        'space_before': 0,
        'space_after': 0,
        'alignment': WD_ALIGN_PARAGRAPH.CENTER,
        'first_line_indent': 0,
        'left_indent': 0,
        'bold': False,
    },
}

# 序号模式（静态文本，禁止自动编号）
# 注意：level1 是居中大标题，通常无序号
NUMBERING_PATTERNS = {
    'level2': r'^[一二三四五六七八九十]+、',           # 一、（二级标题）
    'level3': r'^（[一二三四五六七八九十]+）',         # （一）（三级标题）
    'level4': r'^\d+．',                               # 1．（四级标题，全角点）
    'level5': r'^（\d+）',                             # （1）（五级标题）
    'level6': r'^[①②③④⑤⑥⑦⑧⑨⑩]+',                 # ① 圆圈序号（正文列表项）
}


# ==================== 字体检测与兼容性处理 ====================

def get_available_fonts():
    """获取系统可用字体列表（Windows）"""
    available = set()
    try:
        import winreg
        key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, 
                             r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts")
        i = 0
        while True:
            try:
                name, _, _ = winreg.EnumValue(key, i)
                # 提取字体名称（去除扩展名等信息）
                font_name = name.split(' (')[0].strip()
                available.add(font_name)
                i += 1
            except WindowsError:
                break
        winreg.CloseKey(key)
    except Exception as e:
        print(f"[警告] 无法读取系统字体列表: {e}")
    return available


def get_fallback_font(font_name, available_fonts):
    """获取可用的退化字体"""
    # 先检查原字体是否可用
    if font_name in available_fonts:
        return font_name
    
    # 尝试退化方案
    if font_name in FONT_FALLBACK:
        for fallback in FONT_FALLBACK[font_name]:
            if fallback in available_fonts:
                print(f"[字体退化] {font_name} -> {fallback}")
                return fallback
    
    # 最终退化到系统默认
    print(f"[警告] 字体 {font_name} 不可用，使用系统默认")
    return font_name  # 让 Word 自行处理


class FontChecker:
    """字体兼容性检测器"""
    
    def __init__(self):
        self.available_fonts = get_available_fonts()
        self.font_cache = {}
        self._check_required_fonts()
    
    def _check_required_fonts(self):
        """检查必需字体并输出报告"""
        required = [
            '方正小标宋简体', '黑体', '楷体_GB2312', 
            '仿宋_GB2312', 'Times New Roman'
        ]
        missing = []
        for font in required:
            resolved = self.get_font(font)
            if resolved != font:
                missing.append(f"  {font} -> {resolved}")
        
        if missing:
            print("[字体兼容性报告]")
            print("以下字体进行了退化替换：")
            for m in missing:
                print(m)
        else:
            print("[字体检测] 所有必需字体均可用")
    
    def get_font(self, font_name):
        """获取字体（带缓存）"""
        if font_name not in self.font_cache:
            self.font_cache[font_name] = get_fallback_font(
                font_name, self.available_fonts
            )
        return self.font_cache[font_name]


# ==================== 核心排版引擎 ====================

class GongwenFormatter:
    """公文格式化引擎"""
    
    def __init__(self, doc_path=None):
        """
        初始化格式化引擎
        
        Args:
            doc_path: 可选，已有文档路径。为空则创建新文档
        """
        if doc_path:
            self.doc = Document(doc_path)
        else:
            self.doc = Document()
        
        self.font_checker = FontChecker()
        self._setup_page()
        self._setup_styles()
    
    def _setup_page(self):
        """设置页面参数"""
        section = self.doc.sections[0]
        
        # 纸张大小：A4
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        
        # 页边距
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        
        # 页眉页脚距离
        section.header_distance = HEADER_DISTANCE
        section.footer_distance = FOOTER_DISTANCE
    
    def _setup_styles(self):
        """设置样式库"""
        styles = self.doc.styles
        
        # 设置默认段落样式
        normal = styles['Normal']
        normal.font.name = self.font_checker.get_font('仿宋_GB2312')
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), 
                                        self.font_checker.get_font('仿宋_GB2312'))
    
    def _set_run_font(self, run, cn_font, en_font, font_size, bold=False):
        """设置 run 的字体属性"""
        run.font.name = self.font_checker.get_font(en_font)
        run.font.size = FONT_SIZE.get(font_size, font_size)
        run.font.bold = bold
        
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 
                                     self.font_checker.get_font(cn_font))
    
    def _set_paragraph_format(self, para, style_name):
        """设置段落格式"""
        style = STYLES.get(style_name, STYLES['body'])
        pf = para.paragraph_format
        
        # 对齐方式
        pf.alignment = style['alignment']
        
        # 行距
        if style['line_spacing']:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = style['line_spacing']
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # 段前段后（转换为磅值）
        base_line = Pt(16)  # 三号字约16磅作为基准
        if style['space_before']:
            pf.space_before = Pt(style['space_before'] * 16)
        else:
            pf.space_before = Pt(0)
        
        if style['space_after']:
            pf.space_after = Pt(style['space_after'] * 16)
        else:
            pf.space_after = Pt(0)
        
        # 缩进（字符转磅值，三号字约16磅）
        char_width = FONT_SIZE.get(style['font_size'], Pt(16))
        
        if style['first_line_indent']:
            pf.first_line_indent = Pt(style['first_line_indent'] * 16)
        
        if style['left_indent']:
            pf.left_indent = Pt(style['left_indent'] * 16)
    
    def _process_quotes(self, text, cn_font):
        """
        处理双引号，确保与所在段落中文字体一致
        返回 [(text, is_quote), ...] 列表
        """
        # 匹配中文双引号内容
        pattern = r'(".*?")'
        parts = re.split(pattern, text)
        result = []
        for part in parts:
            if part:
                is_quote = bool(re.match(r'^".*?"$', part))
                result.append((part, is_quote))
        return result
    
    def detect_level(self, text):
        """检测段落级别"""
        text = text.strip()
        
        for level, pattern in NUMBERING_PATTERNS.items():
            if re.match(pattern, text):
                return level
        
        return 'body'
    
    def format_paragraph(self, para, style_name=None):
        """
        格式化单个段落
        
        Args:
            para: 段落对象
            style_name: 指定样式名，为空则自动检测
        """
        text = para.text.strip()
        if not text:
            return
        
        # 自动检测级别
        if style_name is None:
            style_name = self.detect_level(text)
        
        style = STYLES.get(style_name, STYLES['body'])
        cn_font = style['cn_font']
        en_font = style['en_font']
        font_size = style['font_size']
        bold = style.get('bold', False)
        
        # 清除原有 runs
        para.clear()
        
        # 处理双引号逻辑
        parts = self._process_quotes(text, cn_font)
        
        for part_text, is_quote in parts:
            # 分离中英文字符
            segments = self._split_cn_en(part_text)
            for seg_text, is_cn in segments:
                run = para.add_run(seg_text)
                if is_cn:
                    self._set_run_font(run, cn_font, cn_font, font_size, bold)
                else:
                    self._set_run_font(run, cn_font, en_font, font_size, bold)
        
        # 设置段落格式
        self._set_paragraph_format(para, style_name)
    
    def _split_cn_en(self, text):
        """
        分离中英文字符
        返回 [(text, is_chinese), ...] 列表
        """
        result = []
        current_text = ""
        current_is_cn = None
        
        for char in text:
            # 判断是否为中文字符
            is_cn = '\u4e00' <= char <= '\u9fff' or char in '，。、；：""''（）【】《》？！'
            
            if current_is_cn is None:
                current_is_cn = is_cn
                current_text = char
            elif is_cn == current_is_cn:
                current_text += char
            else:
                if current_text:
                    result.append((current_text, current_is_cn))
                current_text = char
                current_is_cn = is_cn
        
        if current_text:
            result.append((current_text, current_is_cn))
        
        return result
    
    def add_heading(self, text, level):
        """
        添加标题
        
        Args:
            text: 标题文本（包含序号）
            level: 级别 (1-5)
        """
        style_name = f'level{level}'
        para = self.doc.add_paragraph()
        
        # 添加文本
        style = STYLES.get(style_name, STYLES['body'])
        cn_font = style['cn_font']
        en_font = style['en_font']
        font_size = style['font_size']
        bold = style.get('bold', False)
        
        # 处理双引号和中英文分离
        parts = self._process_quotes(text, cn_font)
        for part_text, _ in parts:
            segments = self._split_cn_en(part_text)
            for seg_text, is_cn in segments:
                run = para.add_run(seg_text)
                if is_cn:
                    self._set_run_font(run, cn_font, cn_font, font_size, bold)
                else:
                    self._set_run_font(run, cn_font, en_font, font_size, bold)
        
        self._set_paragraph_format(para, style_name)
        return para
    
    def add_body_paragraph(self, text):
        """添加正文段落"""
        para = self.doc.add_paragraph()
        
        style = STYLES['body']
        cn_font = style['cn_font']
        en_font = style['en_font']
        font_size = style['font_size']
        
        # 处理双引号和中英文分离
        parts = self._process_quotes(text, cn_font)
        for part_text, _ in parts:
            segments = self._split_cn_en(part_text)
            for seg_text, is_cn in segments:
                run = para.add_run(seg_text)
                if is_cn:
                    self._set_run_font(run, cn_font, cn_font, font_size)
                else:
                    self._set_run_font(run, cn_font, en_font, font_size)
        
        self._set_paragraph_format(para, 'body')
        return para
    
    def add_table(self, title, rows, cols, data=None, has_header=True):
        """
        添加表格
        
        Args:
            title: 表格标题
            rows: 行数
            cols: 列数
            data: 表格数据 [[row1], [row2], ...]
            has_header: 第一行是否为表头(黑体小四不加粗)
        """
        # 表格标题(表上方空一行)
        title_para = self.doc.add_paragraph()
        style = STYLES['table_title']
        run = title_para.add_run(title)
        self._set_run_font(run, style['cn_font'], style['en_font'], 
                          style['font_size'], style.get('bold', False))
        self._set_paragraph_format(title_para, 'table_title')
        
        # 创建表格
        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格样式
        table.style = 'Table Grid'
        
        # 填充数据并设置格式
        if data:
            for i, row_data in enumerate(data):
                # 判断是否为表头行
                is_header_row = (i == 0 and has_header)
                style_name = 'table_header' if is_header_row else 'table_content'
                style = STYLES[style_name]
                
                for j, cell_text in enumerate(row_data):
                    cell = table.rows[i].cells[j]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    
                    # 分离中英文
                    segments = self._split_cn_en(str(cell_text))
                    for seg_text, is_cn in segments:
                        run = para.add_run(seg_text)
                        if is_cn:
                            self._set_run_font(run, style['cn_font'], 
                                             style['cn_font'], style['font_size'], 
                                             style.get('bold', False))
                        else:
                            self._set_run_font(run, style['cn_font'],
                                             style['en_font'], style['font_size'],
                                             style.get('bold', False))
                    
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 设置单倍行距
                    if style['line_spacing'] is None:
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # 表格后空一行
        self.doc.add_paragraph()
        
        return table
    
    def add_figure_title(self, title, add_spacing=True):
        """
        添加插图标题（位于图下方）
        
        Args:
            title: 插图标题
            add_spacing: 是否在标题后自动添加空行(默认True)
        """
        para = self.doc.add_paragraph()
        style = STYLES['figure_title']
        run = para.add_run(title)
        self._set_run_font(run, style['cn_font'], style['en_font'],
                          style['font_size'], style.get('bold', False))
        self._set_paragraph_format(para, 'figure_title')
        
        # 插图后空一行
        if add_spacing:
            self.doc.add_paragraph()
        
        return para
    
    def add_page_number(self):
        """添加页码（宋体四号，居中，页脚，格式：— 1 —）"""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # 清空页脚原有内容
        for para in footer.paragraphs:
            para.clear()
        
        # 获取或创建段落
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加前导破折号 "— "
        run_prefix = para.add_run('— ')
        run_prefix.font.name = '宋体'
        run_prefix.font.size = FONT_SIZE['四号']
        run_prefix._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加页码域
        run = para.add_run()
        run.font.name = '宋体'
        run.font.size = FONT_SIZE['四号']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 构建页码域代码
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # 添加默认页码文本(1)
        t = OxmlElement('w:t')
        t.text = '1'
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(t)
        run._r.append(fldChar3)
        
        # 添加后置破折号 " —"
        run_suffix = para.add_run(' —')
        run_suffix.font.name = '宋体'
        run_suffix.font.size = FONT_SIZE['四号']
        run_suffix._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def add_header(self, header_text):
        """
        添加页眉
        
        Args:
            header_text: 页眉文字
        """
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        
        # 清空原有内容
        for para in header.paragraphs:
            para.clear()
        
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        
        # 页眉使用仿宋_GB2312 + Times New Roman
        cn_font = '仿宋_GB2312'
        en_font = 'Times New Roman'
        font_size = '三号'
        
        # 处理双引号和中英文分离
        parts = self._process_quotes(header_text, cn_font)
        for part_text, _ in parts:
            segments = self._split_cn_en(part_text)
            for seg_text, is_cn in segments:
                run = para.add_run(seg_text)
                if is_cn:
                    self._set_run_font(run, cn_font, cn_font, font_size)
                else:
                    self._set_run_font(run, cn_font, en_font, font_size)
        
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def format_document(self):
        """格式化整个文档"""
        for para in self.doc.paragraphs:
            self.format_paragraph(para)
    
    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        print(f"[保存成功] {output_path}")


# ==================== Demo 文档生成 ====================

def create_demo_document(output_path):
    """创建包含五级标题、表格、插图的示例文档"""
    
    formatter = GongwenFormatter()
    
    # 添加页码
    formatter.add_page_number()
    
    # 添加页眉
    formatter.add_header('公文格式化示例文档 "Format Demo" 123')
    
    # 一级标题
    formatter.add_heading('一级"标题"123', 1)
    
    # 正文段落
    formatter.add_body_paragraph(
        '这是正文内容示例，采用仿宋_GB2312字体，三号字，'
        '固定28磅行距，首行缩进2字符，两端对齐。'
        '正文中的"引用内容"将保持与段落字体一致。'
        '英文和数字ABC123使用Times New Roman字体。'
    )
    
    # 二级标题
    formatter.add_heading('一、二级"标题"123', 2)
    
    formatter.add_body_paragraph(
        '二级标题采用黑体，三号字，固定30磅行距，左缩进2字符。'
        '序号采用静态文本"一、"格式，禁止使用Word自动编号。'
    )
    
    # 三级标题
    formatter.add_heading('（一）三级"标题"123', 3)
    
    formatter.add_body_paragraph(
        '三级标题采用楷体_GB2312，三号字，固定30磅行距。'
        '序号格式为全角括号"（一）"。'
    )
    
    # 四级标题
    formatter.add_heading('1．四级"标题"123', 4)
    
    formatter.add_body_paragraph(
        '四级标题采用仿宋_GB2312，三号字，固定28磅行距。'
        '注意序号后必须使用全角点"．"而非半角点"."。'
    )
    
    # 五级标题
    formatter.add_heading('（1）五级"标题"123', 5)
    
    formatter.add_body_paragraph(
        '五级标题同样采用仿宋_GB2312格式，序号为"（1）"。'
    )
    
    # 添加带圆圈数字的内容
    formatter.add_body_paragraph(
        '①圆圈数字序号示例：第一点内容说明。'
    )
    formatter.add_body_paragraph(
        '②圆圈数字序号示例：第二点内容说明。'
    )
    
    # 添加表格示例
    table_data = [
        ['序号', '项目名称', '数值', '备注'],
        ['1', '测试项目A', '100', '正常'],
        ['2', '测试项目B', '200', '正常'],
        ['3', '测试项目C', '300', '待确认'],
    ]
    formatter.add_table('表1  示例数据统计表', 4, 4, table_data)
    
    # 添加正文说明
    formatter.add_body_paragraph(
        '上表展示了表格的标准格式：表格标题采用黑体小四，不加粗，居中；'
        '表格内容采用仿宋_GB2312小四，数字使用Times New Roman小四，单倍行距。'
    )
    
    # 添加插图标题示例（实际图片需要用户插入）
    formatter.add_body_paragraph(
        '（此处可插入图片）'
    )
    formatter.add_figure_title('图1  示例流程图')
    
    formatter.add_body_paragraph(
        '插图标题采用黑体小四，不加粗，居中，位于图下方。'
    )
    
    # 保存文档
    formatter.save(output_path)
    
    return formatter


def format_existing_document(input_path, output_path):
    """格式化已有文档"""
    formatter = GongwenFormatter(input_path)
    formatter.format_document()
    formatter.add_page_number()
    formatter.save(output_path)
    return formatter


# ==================== 命令行接口 ====================

if __name__ == '__main__':
    import sys
    
    print("=" * 60)
    print("公文级 Word 文档自动化排版引擎 v1.0")
    print("=" * 60)
    
    if len(sys.argv) < 2:
        # 默认生成 Demo 文档
        output_path = os.path.join(os.path.dirname(__file__), 'demo_gongwen.docx')
        print(f"\n[模式] 生成示例文档")
        create_demo_document(output_path)
    else:
        input_path = sys.argv[1]
        output_path = sys.argv[2] if len(sys.argv) > 2 else input_path.replace('.docx', '_formatted.docx')
        print(f"\n[模式] 格式化已有文档")
        print(f"[输入] {input_path}")
        print(f"[输出] {output_path}")
        format_existing_document(input_path, output_path)
    
    print("\n[完成] 排版引擎执行结束")
